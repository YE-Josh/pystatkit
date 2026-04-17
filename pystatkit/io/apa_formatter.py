"""APA 7 formatting for statistical output (v0.2).

Handles all method types implemented in v0.2:
- demographic: flat table (already multi-column from tableone)
- two_group / paired / one_way_anova: primary + descriptives + post-hoc
- anova_rm: primary (with sphericity info) + post-hoc
- anova_mixed: primary + simple effects + post-hoc
- correlation: pairwise table + correlation matrix
- ancova: primary + adjusted means + assumption check results

The formatter reads method-specific extras from `result.extras` and
appends the right auxiliary tables.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

from pystatkit.core.provenance import RunMetadata
from pystatkit.core.results import AnalysisResult


def format_p_value(p: float) -> str:
    """APA: 'p < .001' or 'p = .03' (no leading zero)."""
    if p < 0.001:
        return "< .001"
    return f"= {p:.3f}".replace("0.", ".")


def _round_df(df: pd.DataFrame, decimals: int = 3) -> pd.DataFrame:
    out = df.copy()
    for col in out.columns:
        if pd.api.types.is_numeric_dtype(out[col]):
            out[col] = out[col].round(decimals)
    return out


def _df_to_rows(df: pd.DataFrame) -> tuple[list[str], list[list[str]]]:
    """Turn a DataFrame into (header, rows of strings) for docx output.

    Handles DataFrames that may or may not have a meaningful index.
    """
    out = df.copy()
    if out.index.name or not isinstance(out.index, pd.RangeIndex):
        out = out.reset_index()
    out = _round_df(out)
    headers = [str(c) for c in out.columns]
    rows = [[str(v) for v in row] for row in out.itertuples(index=False)]
    return headers, rows


# =============================================================================
# XLSX writer — one workbook per method with a sheet per sub-table
# =============================================================================

def write_xlsx_report(
    result: AnalysisResult,
    metadata: RunMetadata,
    output_path: str | Path,
) -> Path:
    """Write a multi-sheet .xlsx containing all tables for one method."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        # Summary sheet.
        summary_rows = [
            ("Method", result.method),
            ("n", result.n_total),
            ("Interpretation", result.interpretation),
            ("", ""),
            ("--- Effect sizes ---", ""),
            *[(k, v) for k, v in result.effect_size.items()],
            ("", ""),
            ("--- Run metadata ---", ""),
            *[(k, str(v)) for k, v in metadata.to_dict().items()],
        ]
        pd.DataFrame(summary_rows, columns=["Field", "Value"]).to_excel(
            writer, sheet_name="Summary", index=False
        )

        if result.descriptives is not None:
            _round_df(result.descriptives).to_excel(
                writer, sheet_name="Descriptives"
            )

        # Primary table.
        primary_out = _round_df(
            result.primary if isinstance(result.primary, pd.DataFrame)
            else pd.DataFrame(result.primary)
        )
        primary_out.to_excel(writer, sheet_name="Primary", index=False)

        if result.posthoc is not None:
            _round_df(result.posthoc).to_excel(
                writer, sheet_name="PostHoc", index=False
            )

        # Method-specific extras.
        _write_extras_to_xlsx(result, writer)

    return output_path


def _write_extras_to_xlsx(result: AnalysisResult, writer) -> None:
    """Append method-specific auxiliary tables as extra sheets."""
    ex = result.extras

    # Correlation matrix.
    matrix = ex.get("matrix")
    if matrix is not None and isinstance(matrix, pd.DataFrame):
        _round_df(matrix).to_excel(writer, sheet_name="CorrelationMatrix")

    # ANCOVA adjusted means.
    adj = ex.get("adjusted_means")
    if adj is not None and isinstance(adj, pd.DataFrame):
        _round_df(adj).to_excel(writer, sheet_name="AdjustedMeans", index=False)

    # ANCOVA assumptions (one sheet per assumption check).
    assumptions = ex.get("assumptions")
    if isinstance(assumptions, dict):
        for name, df_a in assumptions.items():
            if isinstance(df_a, pd.DataFrame):
                sheet = f"Assume_{name}"[:31]  # Excel sheet name limit
                _round_df(df_a).to_excel(writer, sheet_name=sheet, index=False)

    # Mixed ANOVA simple effects.
    simple = ex.get("simple_effects")
    if simple is not None and isinstance(simple, pd.DataFrame):
        _round_df(simple).to_excel(writer, sheet_name="SimpleEffects", index=False)

    # RM-ANOVA sphericity info.
    sph = ex.get("sphericity")
    if isinstance(sph, dict) and sph:
        pd.DataFrame(
            [(k, str(v)) for k, v in sph.items()], columns=["Field", "Value"]
        ).to_excel(writer, sheet_name="Sphericity", index=False)


# =============================================================================
# DOCX writer — APA-styled report for one method
# =============================================================================

def _add_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    """Append a titled table rendered from a DataFrame."""
    doc.add_heading(title, level=2)
    headers, rows = _df_to_rows(df)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, col in enumerate(headers):
        hdr[i].text = col
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = val


def write_docx_report(
    result: AnalysisResult,
    metadata: RunMetadata,
    output_path: str | Path,
) -> Path:
    """APA-styled .docx report for one method."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    doc.add_heading(f"{result.method}", level=1)

    # APA-ready interpretation sentence.
    doc.add_heading("Result", level=2)
    p = doc.add_paragraph(result.interpretation)
    p.paragraph_format.space_after = Pt(12)

    # Descriptives (omit for demographic — its primary table already IS descriptives).
    if result.descriptives is not None and result.method_key != "demographic":
        _add_table(doc, result.descriptives, "Descriptive Statistics")

    # Primary table (heading varies by method for clarity).
    primary_title = _primary_heading(result.method_key)
    _add_table(doc, result.primary, primary_title)

    # Post-hoc.
    if result.posthoc is not None:
        ph_method = result.extras.get("posthoc_method", "post-hoc")
        _add_table(doc, result.posthoc, f"Post-hoc Comparisons ({ph_method})")

    # Method-specific extras rendered as additional tables.
    _add_extras_to_docx(doc, result)

    # Effect size summary (for methods that care).
    if result.effect_size:
        doc.add_heading("Effect Size(s)", level=2)
        for name, val in result.effect_size.items():
            doc.add_paragraph(f"{name}: {val:.3f}", style="List Bullet")

    # Provenance.
    doc.add_heading("Run Metadata", level=2)
    for k, v in metadata.to_dict().items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    doc.save(output_path)
    return output_path


def _primary_heading(method_key: str) -> str:
    """Per-method heading label for the primary table."""
    headings = {
        "demographic": "Demographic / Baseline Characteristics",
        "correlation": "Pairwise Correlations",
        "anova_rm": "RM-ANOVA Results (with Sphericity)",
        "anova_mixed": "Mixed ANOVA Results",
        "ancova": "ANCOVA Results",
    }
    return headings.get(method_key, "Primary Test Statistics")


def _add_extras_to_docx(doc: Document, result: AnalysisResult) -> None:
    """Append method-specific auxiliary tables."""
    ex = result.extras

    # Correlation matrix.
    matrix = ex.get("matrix")
    if matrix is not None and isinstance(matrix, pd.DataFrame):
        _add_table(doc, matrix, "Correlation Matrix")

    # ANCOVA adjusted means.
    adj = ex.get("adjusted_means")
    if adj is not None and isinstance(adj, pd.DataFrame):
        _add_table(doc, adj, "Adjusted (Estimated Marginal) Means")

    # ANCOVA assumption checks.
    assumptions = ex.get("assumptions")
    if isinstance(assumptions, dict) and assumptions:
        doc.add_heading("Assumption Checks", level=2)
        for name, df_a in assumptions.items():
            if isinstance(df_a, pd.DataFrame) and not df_a.empty:
                label = name.replace("_", " ").title()
                _add_table(doc, df_a, label)

    # RM-ANOVA sphericity.
    sph = ex.get("sphericity")
    if isinstance(sph, dict) and sph:
        doc.add_heading("Sphericity (Mauchly's Test)", level=2)
        for k, v in sph.items():
            doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    # Mixed ANOVA simple effects.
    simple = ex.get("simple_effects")
    if simple is not None and isinstance(simple, pd.DataFrame) and not simple.empty:
        _add_table(doc, simple, "Simple Effects (at each between-group level)")
